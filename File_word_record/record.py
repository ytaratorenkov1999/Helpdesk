from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from Result.total import total_request


def record_file(sp):
    # Запись в Word-документ всех заявок
    doc = Document()
    table = doc.add_table(1, len(sp[0]))
    table.style = 'Table Grid'
    head_cells = table.rows[0].cells
    for i, item in enumerate(['Номер заявки', 'Дата и время регистрации обращения', 'Дата и время реакции на обращение',
                              'Дата и время решения обращения', 'Краткое описание обращения',
                              'Сведения об инициаторе обращения (табельный номер/ФИО и подразделение)',
                              'Сведения о ходе решения обращения',
                              'ФИО исполнителя', 'Плановый срок обработки обращения(чаcов)',
                              'Фактический срок обработки обращения(часов)',
                              'Отклонение фактического срока от планового(часов)']):
        p = head_cells[i].paragraphs[0]
        p.add_run(item).bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row in sp:
        cells = table.add_row().cells
        for i, item in enumerate(row):
            cells[i].text = str(item)
    doc.save('Отчет_по_заявкам_Алина.docx')

record_file(total_request())